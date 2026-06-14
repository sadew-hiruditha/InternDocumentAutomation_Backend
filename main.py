from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from docx import Document
import pandas as pd
import io
import zipfile

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Manual Input සඳහා ආකෘතිය
import os

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
NDA_TEMPLATE_PATH = os.path.join(BASE_DIR, "NDA_template.docx")
OFFER_LETTER_TEMPLATE_PATH = os.path.join(BASE_DIR, "Offer_Letter_template.docx")


class InternData(BaseModel):
    name_with_initials: str
    full_name: str
    home_address: str
    welcome_name: str
    start_date: str
    end_date: str
    nic: str
    department: str
    telephone_number: str
    supervisor_name: str
    supervisor_designation: str
    line_address: str


def format_standard_date(date_val):
    if pd.isna(date_val) or not str(date_val).strip():
        return ""
    try:
        dt = pd.to_datetime(date_val)
        return f"{dt.day} {dt.strftime('%B')} {dt.year}"
    except Exception:
        return str(date_val).strip().split(" ")[0]


def format_ordinal_date(date_val):
    if pd.isna(date_val) or not str(date_val).strip():
        return ""
    try:
        dt = pd.to_datetime(date_val)
        day = dt.day
        if 11 <= day <= 13:
            suffix = "th"
        else:
            suffix = {1: "st", 2: "nd", 3: "rd"}.get(day % 10, "th")
        day_str = f"{day}{suffix}"
        month_str = dt.strftime("%B")
        year_str = dt.strftime("%Y")
        return f"{day_str} day of {month_str} {year_str}"
    except Exception:
        return str(date_val).strip()


def normalize_data(raw_data):
    def get_val(keys):
        for k in keys:
            if k in raw_data:
                val = raw_data[k]
                return clean_string_val(val)
            for rk in raw_data.keys():
                rk_clean = str(rk).strip().lower().replace("_", "").replace(" ", "").replace("-", "")
                k_clean = str(k).strip().lower().replace("_", "").replace(" ", "").replace("-", "")
                if rk_clean == k_clean:
                    val = raw_data[rk]
                    return clean_string_val(val)
        return ""

    def clean_string_val(val):
        if pd.isna(val):
            return ""
        if isinstance(val, float) and val.is_integer():
            return str(int(val))
        return str(val).strip()

    return {
        "name_with_initials": get_val(["Name with initials", "Name_with_initials", "name_with_initials", "NameWithInitials", "name"]),
        "full_name": get_val(["Full Name", "Full_Name", "full_name", "FullName"]),
        "home_address": get_val(["Home Adress", "Home Address", "home_address", "Home_Adress", "address"]),
        "welcome_name": get_val(["welcome name", "welcome_name", "welcomeName"]),
        "start_date": get_val(["Start date", "Start_Date", "start_date", "StartDate"]),
        "end_date": get_val(["End date", "End_Date", "end_date", "EndDate"]),
        "nic": get_val(["NIC", "nic", "NIC Number", "nic_number"]),
        "department": get_val(["Department", "department"]),
        "telephone_number": get_val(["Telephone number", "Telephone_number", "telephone_number", "TelephoneNumber", "tel"]),
        "supervisor_name": get_val(["Supervisor name", "Supervisor_name", "supervisor_name", "SupervisorName"]),
        "supervisor_designation": get_val(["supervisor designation", "supervisor_designation", "supervisorDesignation", "Supervisor_Designation"]),
        "line_address": get_val(["line adress", "line address", "line_address", "Line_Adress", "Line_Address"])
    }


def fill_document(template_path, data):
    doc = Document(template_path)
    
    placeholders = {
        "NAME_WITH_INITIALS": data["name_with_initials"],
        "FULL_NAME": data["full_name"],
        "HOME_ADDRESS": data["home_address"],
        "WELCOME_NAME": data["welcome_name"],
        "START_DATE": format_standard_date(data["start_date"]),
        "START_DATE_ORDINAL": format_ordinal_date(data["start_date"]),
        "END_DATE": format_standard_date(data["end_date"]),
        "NIC": data["nic"],
        "DEPARTMENT": data["department"],
        "TELEPHONE": data["telephone_number"],
        "SUPERVISOR_NAME": data["supervisor_name"],
        "SUPERVISOR_DESIGNATION": data["supervisor_designation"],
        "LINE_ADDRESS": data["line_address"]
    }
    
    for p in doc.paragraphs:
        for run in p.runs:
            if run.text:
                for k, v in placeholders.items():
                    tag = f"{{{{{k}}}}}"
                    if tag in run.text:
                        run.text = run.text.replace(tag, str(v))
                        
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        if run.text:
                            for k, v in placeholders.items():
                                tag = f"{{{{{k}}}}}"
                                if tag in run.text:
                                    run.text = run.text.replace(tag, str(v))
                                    
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream.getvalue()


def create_zip_archive(data_list):
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, "a", zipfile.ZIP_DEFLATED, False) as zip_file:
        for raw_data in data_list:
            data = normalize_data(raw_data)
            if not data['full_name'] or data['full_name'] == 'nan':
                continue

            safe_name = data['full_name'].replace(" ", "_")
            ol_file = fill_document(OFFER_LETTER_TEMPLATE_PATH, data)
            nda_file = fill_document(NDA_TEMPLATE_PATH, data)

            zip_file.writestr(f"Offer_Letters/{safe_name}_Offer_Letter.docx", ol_file)
            zip_file.writestr(f"NDAs/{safe_name}_NDA.docx", nda_file)

    zip_buffer.seek(0)
    return zip_buffer


@app.post("/generate")
async def generate_single(data: InternData):
    try:
        zip_buffer = create_zip_archive([data.dict()])
        safe_name = data.full_name.replace(' ', '_')
        return StreamingResponse(
            zip_buffer,
            media_type="application/x-zip-compressed",
            headers={
                "Content-Disposition": f"attachment; filename=Intern_Docs_{safe_name}.zip"}
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/generate-bulk")
async def generate_bulk(file: UploadFile = File(...)):
    if not file.filename.endswith(('.xlsx', '.xls')):
        raise HTTPException(
            status_code=400, detail="කරුණාකර Excel (.xlsx) ෆයිල් එකක් පමණක් Upload කරන්න.")

    try:
        contents = await file.read()
        df = pd.read_excel(io.BytesIO(contents))
        data_list = df.to_dict(orient='records')
        zip_buffer = create_zip_archive(data_list)
        return StreamingResponse(
            zip_buffer,
            media_type="application/x-zip-compressed",
            headers={
                "Content-Disposition": "attachment; filename=Bulk_Intern_Docs.zip"}
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8080)
